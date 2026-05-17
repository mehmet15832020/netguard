"""
BSM498_NetGuard_Raporu.md → profesyonel .docx dönüştürücü
Araştırma temelli en iyi uygulamalar:
  - w:docDefaults ile belge geneli font
  - Heading theme font override (w:asciiTheme temizleme)
  - w:cs ile Türkçe karakter desteği
  - Profesyonel tablo (kenarlık + şerit + hücre padding)
  - Paragraf seviyesi kod bloğu arka plan gölgesi
  - Sayfa numarası (footer, w:fldChar)
  - Kapak sayfası ayrı section
  - 1.5 satır aralığı, hizalı gövde metni
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT       = "Times New Roman"
CODE_FONT  = "Consolas"
BODY_PT    = 12.0
CODE_PT    = 9.5

# Renkler
H1_COLOR   = RGBColor(0x1F, 0x35, 0x64)   # koyu lacivert
H2_COLOR   = RGBColor(0x2E, 0x4D, 0x7B)   # orta lacivert
H3_COLOR   = RGBColor(0x1A, 0x1A, 0x1A)   # neredeyse siyah
TBL_HEAD   = "2E4D7B"                       # tablo başlık zemin
TBL_ALT    = "EEF2F7"                       # tablo çift satır zemin
CODE_BG    = "F4F4F4"                       # kod bloğu zemin


# ─── BELGE DEFAULTs (en düşük öncelik: hepsini etkiler) ─────────────────────

def _set_doc_defaults(doc):
    """w:docDefaults içinde font ve boyut — her paragraf/run için fallback."""
    styles_el = doc.styles.element
    dd = styles_el.find(qn("w:docDefaults"))
    if dd is None:
        dd = OxmlElement("w:docDefaults")
        styles_el.insert(0, dd)
    rPrD = dd.find(qn("w:rPrDefault"))
    if rPrD is None:
        rPrD = OxmlElement("w:rPrDefault")
        dd.append(rPrD)
    rPr = rPrD.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrD.append(rPr)
    # Tüm dört font özelliği
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts")
        rPr.insert(0, rF)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rF.set(qn(attr), FONT)
    for theme in ("w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"):
        rF.attrib.pop(qn(theme), None)
    # Boyut (half-points)
    hp = str(int(BODY_PT * 2))
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn("w:val"), hp)


# ─── NORMAL STİL ─────────────────────────────────────────────────────────────

def _set_normal_style(doc):
    norm = doc.styles["Normal"]
    norm.font.name = FONT
    norm.font.size = Pt(BODY_PT)
    pf = norm.paragraph_format
    pf.alignment             = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule     = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before          = Pt(0)
    pf.space_after           = Pt(6)
    pf.widow_control         = True
    # XML üzerinden cs + eastAsia
    rPr = norm._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts")
        rPr.insert(0, rF)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rF.set(qn(attr), FONT)
    for theme in ("w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"):
        rF.attrib.pop(qn(theme), None)


# ─── BAŞLIK STİLLERİ (theme override dahil) ──────────────────────────────────

def _configure_headings(doc):
    cfg = {
        "Heading 1": (16, True,  True,  18, 8,  H1_COLOR),
        "Heading 2": (14, True,  False, 14, 6,  H2_COLOR),
        "Heading 3": (12, True,  False, 10, 4,  H3_COLOR),
        "Heading 4": (12, False, False,  8, 3,  H3_COLOR),
    }
    for name, (sz, bold, caps, bef, aft, col) in cfg.items():
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.size  = Pt(sz)
        st.font.bold  = bold
        st.font.all_caps = caps
        if col:
            st.font.color.rgb = col
        pf = st.paragraph_format
        pf.space_before          = Pt(bef)
        pf.space_after           = Pt(aft)
        pf.alignment             = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing_rule     = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.keep_with_next        = True
        pf.page_break_before     = False
        # KRITIK: theme fontlarini temizle
        rPr = st._element.get_or_add_rPr()
        rF = rPr.find(qn("w:rFonts"))
        if rF is None:
            rF = OxmlElement("w:rFonts")
            rPr.insert(0, rF)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rF.set(qn(attr), FONT)
        for theme in ("w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"):
            rF.attrib.pop(qn(theme), None)


# ─── SAYFA DÜZENİ ────────────────────────────────────────────────────────────

def _set_margins(section):
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


# ─── SAYFA NUMARASI ──────────────────────────────────────────────────────────

def _field_run(para, code: str):
    run = para.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {code} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _add_page_numbers(doc):
    """Footer'a 'Sayfa X / Y' ekle (tüm section'lar için)."""
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Sayfa ")
        r.font.name = FONT
        r.font.size = Pt(10)
        _field_run(fp, "PAGE")
        r2 = fp.add_run(" / ")
        r2.font.name = FONT
        r2.font.size = Pt(10)
        _field_run(fp, "NUMPAGES")


# ─── FONT YARDIMCI ───────────────────────────────────────────────────────────

def _set_run_font(run, name=FONT, size=None, bold=False, italic=False,
                  color=None, underline=False):
    run.font.name      = name
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.underline = underline
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    # XML: tüm dört özellik + cs (Türkçe karakter)
    rPr = run._r.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = OxmlElement("w:rFonts")
        rPr.insert(0, rF)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rF.set(qn(attr), name)
    for theme in ("w:asciiTheme", "w:hAnsiTheme", "w:csTheme", "w:eastAsiaTheme"):
        rF.attrib.pop(qn(theme), None)


def _set_para_font_spacing(para, before=Pt(0), after=Pt(6),
                            spacing=WD_LINE_SPACING.ONE_POINT_FIVE,
                            align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = para.paragraph_format
    pf.space_before      = before
    pf.space_after       = after
    pf.line_spacing_rule = spacing
    pf.alignment         = align


# ─── INLINE BİÇİMLENDİRME (**bold**, *italic*, `code`) ───────────────────────

_INLINE = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)")

def _inline(para, text: str, size=Pt(BODY_PT), font=FONT, align=None):
    if align is not None:
        para.alignment = align
    pos = 0
    for m in _INLINE.finditer(text):
        before = text[pos:m.start()]
        if before:
            _set_run_font(para.add_run(before), name=font, size=size)
        if m.group(2):
            _set_run_font(para.add_run(m.group(2)), name=font, size=size, bold=True)
        elif m.group(3):
            _set_run_font(para.add_run(m.group(3)), name=font, size=size, italic=True)
        elif m.group(4):
            _set_run_font(para.add_run(m.group(4)), name=CODE_FONT, size=Pt(CODE_PT))
        pos = m.end()
    tail = text[pos:]
    if tail:
        _set_run_font(para.add_run(tail), name=font, size=size)


# ─── BAŞLIK EKLE ──────────────────────────────────────────────────────────────

def _add_heading(doc, text: str, level: int):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style = style_map.get(level, "Heading 4")
    para = doc.add_paragraph(style=style)
    _inline(para, text, size=Pt([16, 14, 12, 12][min(level - 1, 3)]))
    return para


# ─── GÖVDE PARAGRAF ───────────────────────────────────────────────────────────

def _add_body(doc, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    para = doc.add_paragraph()
    _set_para_font_spacing(para, align=align)
    _inline(para, text, align=align)
    return para


# ─── MADDE İŞARETİ ────────────────────────────────────────────────────────────

def _add_bullet(doc, text: str, level: int = 0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _inline(para, text)
    return para


# ─── KOD BLOĞU ────────────────────────────────────────────────────────────────

def _add_code_block(doc, lines: list[str]):
    for raw in lines:
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before      = Pt(0)
        pf.space_after       = Pt(0)
        pf.left_indent       = Cm(0.3)
        pf.right_indent      = Cm(0.3)
        pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
        # Paragraf seviyesi arka plan (w:pPr > w:shd) — araştırmada doğrulanan yöntem
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_BG)
        _remove_existing(pPr, "w:shd")
        pPr.append(shd)
        # Kenarlık (kutu görünümü)
        pBdr = OxmlElement("w:pBdr")
        for side in ("top", "bottom", "left", "right"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "4")
            el.set(qn("w:color"), "CCCCCC")
            pBdr.append(el)
        _remove_existing(pPr, "w:pBdr")
        pPr.append(pBdr)
        run = para.add_run(raw if raw else " ")
        _set_run_font(run, name=CODE_FONT, size=Pt(CODE_PT))


# ─── TABLO ────────────────────────────────────────────────────────────────────

def _remove_existing(parent, tag: str):
    el = parent.find(qn(tag))
    if el is not None:
        parent.remove(el)


def _cell_background(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    _remove_existing(tcPr, "w:shd")
    tcPr.append(shd)


def _cell_borders(cell, color="AAAAAA", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "start", "end", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    _remove_existing(tcPr, "w:tcBorders")
    tcPr.append(tcBorders)


def _cell_padding(cell, top=60, bottom=60, start=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    _remove_existing(tcPr, "w:tcMar")
    tcPr.append(tcMar)


def _is_sep_row(line: str) -> bool:
    return bool(re.match(r"^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?\s*$", line.strip()))


def _parse_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _add_table(doc, rows: list[list[str]]):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style     = "Table Grid"

    # Tablo genişliği (sayfa içi tam genişlik)
    tbl_el = tbl._tbl
    tblPr  = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    _remove_existing(tblPr, "w:tblW")
    tblPr.append(tblW)

    for ri, row_data in enumerate(rows):
        is_header = (ri == 0)
        is_alt    = (ri % 2 == 0) and not is_header
        for ci, cell_text in enumerate(row_data):
            cell = tbl.cell(ri, ci)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _cell_borders(cell, color="BBBBBB", sz="4")
            _cell_padding(cell)
            if is_header:
                _cell_background(cell, TBL_HEAD)
            elif is_alt:
                _cell_background(cell, TBL_ALT)
            para = cell.paragraphs[0]
            para.clear()
            para.paragraph_format.space_before      = Pt(2)
            para.paragraph_format.space_after       = Pt(2)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
            _inline(para, cell_text, size=Pt(10.5))
            if is_header:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Tablodan sonra boşluk
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ─── KAPAK SAYFASI ────────────────────────────────────────────────────────────

def _centered(doc, text: str, bold=False, size_pt=12.0, before=0.0, after=12.0,
               color=None, italic=False, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before      = Pt(before)
    pf.space_after       = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    txt = text.upper() if caps else text
    run = p.add_run(txt)
    _set_run_font(run, size=Pt(size_pt), bold=bold, italic=italic, color=color)
    return p


# ─── ANA DÖNÜŞTÜRÜCÜ ─────────────────────────────────────────────────────────

def _page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    from docx.enum.text import WD_BREAK
    run.add_break(WD_BREAK.PAGE)


def convert(md_path: str, out_path: str):
    doc = Document()

    # 1. Belge düzeni
    _set_margins(doc.sections[0])

    # 2. Belge geneli font defaults
    _set_doc_defaults(doc)

    # 3. Normal stil
    _set_normal_style(doc)

    # 4. Başlık stilleri
    _configure_headings(doc)

    # 5. İçeriği işle
    lines   = Path(md_path).read_text(encoding="utf-8").splitlines()
    i       = 0
    table   : list[list[str]] = []
    in_code = False
    code_lines: list[str] = []

    def flush_table():
        nonlocal table
        if table:
            _add_table(doc, table)
            table = []

    # Kapak sayfası tespiti: ilk &nbsp; satırına kadar olan her şey kapak
    cover_end = 0
    for idx, ln in enumerate(lines):
        if ln.strip() == "&nbsp;":
            cover_end = idx
            break

    in_cover = True  # kapak bölgesi bayrağı

    while i < len(lines):
        line = lines[i]
        raw  = line.strip()

        # Kapak sonu
        if in_cover and i == cover_end:
            in_cover = False
            _page_break(doc)
            i += 1
            continue

        # Kod bloğu başlangıç/bitiş
        if raw.startswith("```"):
            flush_table()
            if not in_code:
                in_code    = True
                code_lines = []
            else:
                in_code = False
                _add_code_block(doc, code_lines)
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line.rstrip())
            i += 1
            continue

        # Yatay çizgi (-- veya ---+ )
        if re.match(r"^-{3,}$", raw):
            flush_table()
            if not in_cover:
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Sayfa sonu işareti
        if raw == "&nbsp;":
            flush_table()
            _page_break(doc)
            i += 1
            continue

        # Boş satır
        if not raw:
            flush_table()
            i += 1
            continue

        # Tablo satırı
        if raw.startswith("|"):
            if _is_sep_row(raw):
                i += 1
                continue
            table.append(_parse_row(raw))
            i += 1
            continue
        else:
            flush_table()

        # Başlık
        hm = re.match(r"^(#{1,6})\s+(.*)", raw)
        if hm:
            level = len(hm.group(1))
            text  = hm.group(2).strip()
            if in_cover:
                # Kapak: ortalanmış büyük metin
                sz_map = {1: 16.0, 2: 13.0, 3: 12.0}
                sz = sz_map.get(level, 12.0)
                bef = 24.0 if level == 1 else 6.0
                _centered(doc, text, bold=True, size_pt=sz, before=bef, after=6.0,
                          color=H1_COLOR if level == 1 else H2_COLOR)
            else:
                _add_heading(doc, text, level)
            i += 1
            continue

        # Madde işareti
        bm = re.match(r"^(\s*)-\s+(.*)", raw)
        if bm:
            indent = len(bm.group(1)) // 2
            _add_bullet(doc, bm.group(2), level=indent)
            i += 1
            continue

        # Numaralı liste
        nm = re.match(r"^\d+\.\s+(.*)", raw)
        if nm:
            _add_bullet(doc, nm.group(1))
            i += 1
            continue

        # Normal paragraf
        align = WD_ALIGN_PARAGRAPH.CENTER if in_cover else WD_ALIGN_PARAGRAPH.JUSTIFY
        _add_body(doc, raw, align=align)
        i += 1

    flush_table()

    # 6. Sayfa numaraları
    _add_page_numbers(doc)

    doc.save(out_path)
    size_kb = Path(out_path).stat().st_size // 1024
    print(f"✅ Kaydedildi: {out_path}  ({size_kb} KB)")


if __name__ == "__main__":
    base = Path(__file__).parent.parent
    md  = str(base / "BSM498_NetGuard_Raporu.md")
    out = str(base / "BSM498_NetGuard_Raporu.docx")
    convert(md, out)
